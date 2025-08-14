#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
建筑安全分析报告生成器
支持Word、PDF格式导出
"""

import os
import json
from datetime import datetime
from typing import Dict, List, Any

class ReportGenerator:
    """建筑安全分析报告生成器"""
    
    def __init__(self):
        self.company_name = "建筑安全检测平台"
        self.report_template = {
            "title": "建筑安全与质量检测报告",
            "subtitle": "基于AI视觉识别技术的安全分析",
            "footer": "本报告由AI系统自动生成，仅供参考"
        }
    
    def generate_word_report(self, analysis_data: Dict[str, Any], output_path: str) -> bool:
        """生成Word格式报告"""
        try:
            from docx import Document
            from docx.shared import Inches, Pt
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # 创建文档
            doc = Document()
            
            # 标题
            title = doc.add_heading(self.report_template["title"], 0)
            title.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 基本信息
            doc.add_heading("基本信息", level=1)
            info_table = doc.add_table(rows=4, cols=2)
            info_table.style = 'Table Grid'
            
            info_data = [
                ["检测时间", analysis_data.get("timestamp", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))],
                ["检测地点", analysis_data.get("location", "未知")],
                ["检测人员", analysis_data.get("inspector", "AI系统")],
                ["报告编号", analysis_data.get("report_id", "AI-" + datetime.now().strftime("%Y%m%d%H%M%S"))]
            ]
            
            for i, (key, value) in enumerate(info_data):
                info_table.cell(i, 0).text = key
                info_table.cell(i, 1).text = str(value)
            
            # 安全评分
            doc.add_heading("安全评分", level=1)
            score = analysis_data.get("summary", {}).get("total_score", 0)
            doc.add_paragraph(f"整体安全评分：{score}/100")
            
            # 违规统计
            doc.add_heading("违规统计", level=1)
            violations = analysis_data.get("violations", [])
            if violations:
                stats_table = doc.add_table(rows=len(violations) + 1, cols=4)
                stats_table.style = 'Table Grid'
                
                headers = ["序号", "违规类型", "严重程度", "风险等级"]
                for i, header in enumerate(headers):
                    stats_table.cell(0, i).text = header
                
                for i, violation in enumerate(violations):
                    stats_table.cell(i + 1, 0).text = str(i + 1)
                    stats_table.cell(i + 1, 1).text = violation.get("type", "未知")
                    stats_table.cell(i + 1, 2).text = violation.get("severity", "未知")
                    stats_table.cell(i + 1, 3).text = violation.get("risk_level", "未知")
            else:
                doc.add_paragraph("未发现明显违规行为")
            
            # 详细违规信息
            if violations:
                doc.add_heading("详细违规信息", level=1)
                for i, violation in enumerate(violations):
                    doc.add_heading(f"违规 {i + 1}: {violation.get('category', '未知类别')}", level=2)
                    doc.add_paragraph(f"违规行为：{violation.get('description', '无描述')}")
                    
                    regulations = violation.get("regulations", [])
                    if regulations:
                        doc.add_paragraph("违反规范：")
                        for reg in regulations:
                            doc.add_paragraph(f"• {reg.get('code', '')} 第{reg.get('article', '')}条：{reg.get('content', '')}", style='List Bullet')
                    
                    suggestions = violation.get("suggestions", [])
                    if suggestions:
                        doc.add_paragraph("整改建议：")
                        for suggestion in suggestions:
                            doc.add_paragraph(f"• {suggestion}", style='List Bullet')
            
            # 整体评估
            doc.add_heading("整体安全评估", level=1)
            assessment = analysis_data.get("summary", {}).get("overall_assessment", "无评估")
            doc.add_paragraph(assessment)
            
            # 保存文档
            doc.save(output_path)
            return True
            
        except ImportError:
            print("❌ 缺少python-docx库，请运行: pip install python-docx")
            return False
        except Exception as e:
            print(f"❌ Word报告生成失败: {e}")
            return False
    
    def generate_pdf_report(self, analysis_data: Dict[str, Any], output_path: str) -> bool:
        """生成PDF格式报告"""
        try:
            from reportlab.lib.pagesizes import A4
            from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
            from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
            from reportlab.lib.units import inch
            from reportlab.lib import colors
            
            # 创建PDF文档
            doc = SimpleDocTemplate(output_path, pagesize=A4)
            styles = getSampleStyleSheet()
            
            # 自定义样式
            title_style = ParagraphStyle(
                'CustomTitle',
                parent=styles['Heading1'],
                fontSize=18,
                spaceAfter=30,
                alignment=1
            )
            
            # 构建内容
            story = []
            
            # 标题
            story.append(Paragraph(self.report_template["title"], title_style))
            story.append(Spacer(1, 20))
            
            # 基本信息
            story.append(Paragraph("基本信息", styles['Heading1']))
            story.append(Spacer(1, 12))
            
            info_data = [
                ["检测时间", analysis_data.get("timestamp", datetime.now().strftime("%Y-%m-%d %H:%M:%S"))],
                ["检测地点", analysis_data.get("location", "未知")],
                ["检测人员", analysis_data.get("inspector", "AI系统")],
                ["报告编号", analysis_data.get("report_id", "AI-" + datetime.now().strftime("%Y%m%d%H%M%S"))]
            ]
            
            info_table = Table(info_data, colWidths=[2*inch, 4*inch])
            info_table.setStyle(TableStyle([
                ('GRID', (0, 0), (-1, -1), 1, colors.black),
                ('BACKGROUND', (0, 0), (0, -1), colors.grey),
                ('TEXTCOLOR', (0, 0), (0, -1), colors.whitesmoke),
                ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                ('FONTNAME', (0, 0), (-1, -1), 'Helvetica-Bold'),
                ('FONTSIZE', (0, 0), (-1, -1), 10),
                ('BOTTOMPADDING', (0, 0), (-1, -1), 12),
                ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ]))
            story.append(info_table)
            story.append(Spacer(1, 20))
            
            # 安全评分
            story.append(Paragraph("安全评分", styles['Heading1']))
            story.append(Spacer(1, 12))
            score = analysis_data.get("summary", {}).get("total_score", 0)
            story.append(Paragraph(f"整体安全评分：{score}/100", styles['Normal']))
            story.append(Spacer(1, 20))
            
            # 违规统计
            story.append(Paragraph("违规统计", styles['Heading1']))
            story.append(Spacer(1, 12))
            
            violations = analysis_data.get("violations", [])
            if violations:
                stats_data = [["序号", "违规类型", "严重程度", "风险等级"]]
                for i, violation in enumerate(violations):
                    stats_data.append([
                        str(i + 1),
                        violation.get("type", "未知"),
                        violation.get("severity", "未知"),
                        violation.get("risk_level", "未知")
                    ])
                
                stats_table = Table(stats_data, colWidths=[0.8*inch, 2*inch, 1.5*inch, 1.5*inch])
                stats_table.setStyle(TableStyle([
                    ('GRID', (0, 0), (-1, -1), 1, colors.black),
                    ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
                    ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
                    ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
                    ('FONTNAME', (0, 0), (-1, 0), 'Helvetica-Bold'),
                    ('FONTSIZE', (0, 0), (-1, -1), 9),
                ]))
                story.append(stats_table)
                story.append(Spacer(1, 20))
                
                # 详细违规信息
                story.append(Paragraph("详细违规信息", styles['Heading1']))
                story.append(Spacer(1, 12))
                
                for i, violation in enumerate(violations):
                    story.append(Paragraph(f"违规 {i + 1}: {violation.get('category', '未知类别')}", styles['Heading2']))
                    story.append(Paragraph(f"违规行为：{violation.get('description', '无描述')}", styles['Normal']))
                    
                    regulations = violation.get("regulations", [])
                    if regulations:
                        story.append(Paragraph("违反规范：", styles['Normal']))
                        for reg in regulations:
                            story.append(Paragraph(f"• {reg.get('code', '')} 第{reg.get('article', '')}条：{reg.get('content', '')}", styles['Normal']))
                    
                    suggestions = violation.get("suggestions", [])
                    if suggestions:
                        story.append(Paragraph("整改建议：", styles['Normal']))
                        for suggestion in suggestions:
                            story.append(Paragraph(f"• {suggestion}", styles['Normal']))
                    
                    story.append(Spacer(1, 12))
            else:
                story.append(Paragraph("未发现明显违规行为", styles['Normal']))
                story.append(Spacer(1, 20))
            
            # 整体评估
            story.append(Paragraph("整体安全评估", styles['Heading1']))
            story.append(Spacer(1, 12))
            assessment = analysis_data.get("summary", {}).get("overall_assessment", "无评估")
            story.append(Paragraph(assessment, styles['Normal']))
            
            # 生成PDF
            doc.build(story)
            return True
            
        except ImportError:
            print("❌ 缺少reportlab库，请运行: pip install reportlab")
            return False
        except Exception as e:
            print(f"❌ PDF报告生成失败: {e}")
            return False
    
    def generate_report(self, analysis_data: Dict[str, Any], format_type: str = "word", output_dir: str = "./reports") -> str:
        """生成指定格式的报告"""
        # 确保输出目录存在
        os.makedirs(output_dir, exist_ok=True)
        
        # 生成文件名
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"建筑安全报告_{timestamp}"
        
        if format_type.lower() == "word":
            output_path = os.path.join(output_dir, f"{filename}.docx")
            if self.generate_word_report(analysis_data, output_path):
                return output_path
        elif format_type.lower() == "pdf":
            output_path = os.path.join(output_dir, f"{filename}.pdf")
            if self.generate_pdf_report(analysis_data, output_path):
                return output_path
        else:
            print(f"❌ 不支持的报告格式: {format_type}")
            return ""
        
        return ""

# 使用示例
if __name__ == "__main__":
    # 测试数据
    test_data = {
        "timestamp": "2025-08-14 22:30:00",
        "location": "测试建筑工地",
        "inspector": "AI系统",
        "report_id": "AI-20250814223000",
        "violations": [
            {
                "type": "严重违规",
                "category": "安全防护",
                "description": "工人未佩戴安全帽",
                "severity": "high",
                "risk_level": "极高风险",
                "regulations": [
                    {
                        "code": "JGJ59-2011",
                        "article": "3.1.1",
                        "content": "进入施工现场必须佩戴安全帽"
                    }
                ],
                "suggestions": [
                    "立即停止作业",
                    "发放安全帽",
                    "加强安全教育"
                ]
            }
        ],
        "summary": {
            "total_score": 75,
            "overall_assessment": "存在安全隐患，需要立即整改",
            "priority_actions": [
                "立即发放安全防护用品",
                "加强现场安全巡查",
                "开展安全教育培训"
            ]
        }
    }
    
    # 创建报告生成器
    generator = ReportGenerator()
    
    # 生成不同格式的报告
    print("🔍 测试报告生成功能...")
    
    # Word报告
    word_path = generator.generate_report(test_data, "word")
    if word_path:
        print(f"✅ Word报告生成成功: {word_path}")
    
    # PDF报告
    pdf_path = generator.generate_report(test_data, "pdf")
    if pdf_path:
        print(f"✅ PDF报告生成成功: {pdf_path}")
    
    print("🎯 报告生成测试完成！")
