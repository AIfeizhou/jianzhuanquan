# -*- coding: utf-8 -*-
import os
import json
import argparse
from datetime import datetime

def generate_word_report(data, output_path):
    """生成Word格式报告，不包含标注照片"""
    try:
        from docx import Document
        from docx.shared import Inches, Pt
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        doc = Document()
        
        # 设置页面边距
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(1)
            section.bottom_margin = Inches(1)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # 标题
        title = doc.add_heading('建筑安全分析报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 生成时间
        time_para = doc.add_paragraph()
        time_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        time_run = time_para.add_run(f'生成时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}')
        time_run.font.size = Pt(12)
        
        doc.add_paragraph()  # 空行
        
        # 分析概览
        doc.add_heading('分析概览', level=1)
        
        # 创建表格
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # 设置表格样式
        hdr_cells = table.rows[0].cells
        hdr_cells[0].text = '项目'
        hdr_cells[1].text = '数值'
        
        # 添加数据行
        summary = data.get('summary', {})
        table_data = [
            ('安全评分', f"{summary.get('total_score', 0)}"),
            ('严重违规', f"{summary.get('severe_count', 0)}"),
            ('一般违规', f"{summary.get('normal_count', 0)}"),
            ('总违规数', f"{summary.get('severe_count', 0) + summary.get('normal_count', 0)}")
        ]
        
        for item, value in table_data:
            row_cells = table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = str(value)
        
        doc.add_paragraph()  # 空行
        
        # 整体评估
        doc.add_heading('整体评估', level=1)
        assessment = doc.add_paragraph(summary.get('overall_assessment', '未能生成评估报告'))
        
        # 优先整改事项
        doc.add_heading('优先整改事项', level=1)
        priority_actions = summary.get('priority_actions', [])
        if priority_actions:
            for action in priority_actions:
                doc.add_paragraph(f'• {action}', style='List Bullet')
        else:
            doc.add_paragraph('暂无优先整改事项')
        
        doc.add_paragraph()  # 空行
        
        # 违规详情
        violations = data.get('violations', [])
        if violations:
            doc.add_heading('违规详情', level=1)
            
            for i, violation in enumerate(violations, 1):
                # 违规标题
                violation_title = f"{i}. {violation.get('type', '违规')} - {violation.get('category', '建筑安全违规')}"
                doc.add_heading(violation_title, level=2)
                
                # 违规描述
                doc.add_paragraph(f"违规描述: {violation.get('description', '无描述')}")
                
                # 相关条例
                regulations = violation.get('regulations', [])
                if regulations:
                    doc.add_paragraph("相关条例:")
                    for reg in regulations:
                        reg_text = f"• {reg.get('code', '')} {reg.get('article', '')}: {reg.get('content', '')}"
                        doc.add_paragraph(reg_text, style='List Bullet')
                
                # 整改建议
                suggestions = violation.get('suggestions', [])
                if suggestions:
                    doc.add_paragraph("整改建议:")
                    for suggestion in suggestions:
                        doc.add_paragraph(f"• {suggestion}", style='List Bullet')
                
                # 风险等级
                risk_level = violation.get('risk_level', '')
                if risk_level:
                    doc.add_paragraph(f"风险等级: {risk_level}")
                
                doc.add_paragraph()  # 空行
        
        # 保存文档
        doc.save(output_path)
        print(f"Word报告已生成: {output_path}")
        return True
        
    except ImportError:
        print("缺少python-docx库，请运行: pip install python-docx")
        return False
    except Exception as e:
        print(f"Word报告生成失败: {e}")
        return False

def generate_pdf_report(data, output_path):
    """生成PDF格式报告，不包含标注照片"""
    try:
        from reportlab.lib.pagesizes import A4
        from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
        from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
        from reportlab.lib.units import inch
        from reportlab.lib import colors
        
        doc = SimpleDocTemplate(output_path, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        
        # 尝试使用中文字体，如果失败则回退到默认字体
        try:
            from reportlab.pdfbase import pdfmetrics
            from reportlab.pdfbase.ttfonts import TTFont
            
            # 尝试多个中文字体路径
            font_paths = [
                "C:/Windows/Fonts/simsun.ttc",  # Windows宋体
                "C:/Windows/Fonts/msyh.ttc",    # 微软雅黑
                "C:/Windows/Fonts/simhei.ttf",  # 黑体
                "C:/Windows/Fonts/simsun.ttf"   # 宋体TTF
            ]
            
            chinese_font = 'Helvetica'  # 默认字体
            
            for font_path in font_paths:
                if os.path.exists(font_path):
                    try:
                        pdfmetrics.registerFont(TTFont('ChineseFont', font_path))
                        chinese_font = 'ChineseFont'
                        print(f"成功注册字体: {font_path}")
                        break
                    except Exception as e:
                        print(f"字体注册失败 {font_path}: {e}")
                        continue
                        
        except Exception as e:
            print(f"字体处理失败: {e}")
            chinese_font = 'Helvetica'
        
        # 自定义样式
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=18,
            spaceAfter=30,
            alignment=1,  # 居中
            fontName=chinese_font
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12,
            spaceBefore=20,
            fontName=chinese_font
        )
        
        normal_style = ParagraphStyle(
            'CustomNormal',
            parent=styles['Normal'],
            fontSize=10,
            spaceAfter=6,
            fontName=chinese_font
        )
        
        # 标题
        story.append(Paragraph('建筑安全分析报告', title_style))
        story.append(Spacer(1, 20))
        
        # 生成时间
        time_text = f'生成时间: {datetime.now().strftime("%Y年%m月%d日 %H:%M:%S")}'
        story.append(Paragraph(time_text, normal_style))
        story.append(Spacer(1, 20))
        
        # 分析概览
        story.append(Paragraph('分析概览', heading_style))
        
        summary = data.get('summary', {})
        table_data = [
            ['项目', '数值'],
            ['安全评分', str(summary.get('total_score', 0))],
            ['严重违规', str(summary.get('severe_count', 0))],
            ['一般违规', str(summary.get('normal_count', 0))],
            ['总违规数', str(summary.get('severe_count', 0) + summary.get('normal_count', 0))]
        ]
        
        table = Table(table_data, colWidths=[2*inch, 1.5*inch])
        table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (-1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (-1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
            ('FONTNAME', (0, 0), (-1, 0), chinese_font),
            ('FONTSIZE', (0, 0), (-1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (-1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('FONTNAME', (0, 1), (-1, -1), chinese_font),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        
        story.append(table)
        story.append(Spacer(1, 20))
        
        # 整体评估
        story.append(Paragraph('整体评估', heading_style))
        assessment = summary.get('overall_assessment', '未能生成评估报告')
        story.append(Paragraph(assessment, normal_style))
        story.append(Spacer(1, 20))
        
        # 优先整改事项
        story.append(Paragraph('优先整改事项', heading_style))
        priority_actions = summary.get('priority_actions', [])
        if priority_actions:
            for action in priority_actions:
                story.append(Paragraph(f'• {action}', normal_style))
        else:
            story.append(Paragraph('暂无优先整改事项', normal_style))
        
        story.append(Spacer(1, 20))
        
        # 违规详情
        violations = data.get('violations', [])
        if violations:
            story.append(Paragraph('违规详情', heading_style))
            
            for i, violation in enumerate(violations, 1):
                # 违规标题
                violation_title = f"{i}. {violation.get('type', '违规')} - {violation.get('category', '建筑安全违规')}"
                story.append(Paragraph(violation_title, heading_style))
                
                # 违规描述
                description = violation.get('description', '无描述')
                story.append(Paragraph(f"违规描述: {description}", normal_style))
                
                # 相关条例
                regulations = violation.get('regulations', [])
                if regulations:
                    story.append(Paragraph("相关条例:", normal_style))
                    for reg in regulations:
                        reg_text = f"• {reg.get('code', '')} {reg.get('article', '')}: {reg.get('content', '')}"
                        story.append(Paragraph(reg_text, normal_style))
                
                # 整改建议
                suggestions = violation.get('suggestions', [])
                if suggestions:
                    story.append(Paragraph("整改建议:", normal_style))
                    for suggestion in suggestions:
                        story.append(Paragraph(f"• {suggestion}", normal_style))
                
                # 风险等级
                risk_level = violation.get('risk_level', '')
                if risk_level:
                    story.append(Paragraph(f"风险等级: {risk_level}", normal_style))
                
                story.append(Spacer(1, 15))
        
        # 生成PDF
        doc.build(story)
        print(f"PDF报告已生成: {output_path}")
        return True
        
    except ImportError:
        print("缺少reportlab库，请运行: pip install reportlab")
        return False
    except Exception as e:
        print(f"PDF报告生成失败: {e}")
        return False

def generate_report(data, format_type='pdf', output_dir='./temp'):
    """生成指定格式的报告"""
    # 确保输出目录存在
    os.makedirs(output_dir, exist_ok=True)
    
    # 生成文件名
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    
    if format_type == 'word':
        filename = f"Building_Safety_Report_{timestamp}.docx"
        output_path = os.path.join(output_dir, filename)
        return generate_word_report(data, output_path)
    elif format_type == 'pdf':
        filename = f"Building_Safety_Report_{timestamp}.pdf"
        output_path = os.path.join(output_dir, filename)
        return generate_pdf_report(data, output_path)
    else:
        print(f"不支持的格式: {format_type}")
        return False

def main():
    parser = argparse.ArgumentParser(description='生成建筑安全分析报告')
    parser.add_argument('--format', choices=['pdf', 'word'], default='pdf', help='报告格式 (默认: pdf)')
    parser.add_argument('--data', help='分析数据JSON文件路径')
    parser.add_argument('--output', default='./temp', help='输出目录 (默认: ./temp)')
    
    args = parser.parse_args()
    
    # 如果没有提供数据文件，使用示例数据
    if args.data and os.path.exists(args.data):
        try:
            with open(args.data, 'r', encoding='utf-8') as f:
                data = json.load(f)
        except Exception as e:
            print(f"读取数据文件失败: {e}")
            return
    else:
        # 使用示例数据
        data = {
            "violations": [
                {
                    "type": "严重违规",
                    "category": "基坑支护安全",
                    "description": "沟槽深度超过1.5m，两侧边缘未设置标准防护栏杆",
                    "regulations": [{
                        "code": "JGJ59-2011",
                        "article": "4.1.3",
                        "content": "基坑深度超过1.5m时，必须设置安全防护栏杆"
                    }],
                    "suggestions": ["立即设置安全防护栏杆", "加强现场安全巡查"],
                    "risk_level": "极高风险"
                }
            ],
            "summary": {
                "severe_count": 1,
                "normal_count": 0,
                "total_score": 60,
                "overall_assessment": "存在严重安全隐患，需要立即整改",
                "priority_actions": ["立即设置基坑安全防护栏杆"]
            }
        }
    
    # 生成报告
    success = generate_report(data, args.format, args.output)
    
    if success:
        print(f"{args.format.upper()}格式报告生成成功！")
    else:
        print(f"{args.format.upper()}格式报告生成失败！")

if __name__ == '__main__':
    main()
